from docx import Document
from docx.shared import Pt

def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)

def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Header
doc.add_heading('Sudip Bandopadhyay', 0)
doc.add_paragraph(
    'Lead/Senior Data Engineer (GCP/Azure) | Data Scientist\n'
    'Whiteley, UK | +44 7747292241 | sudip.bandopadhyay23@gmail.com | '
    'LinkedIn: linkedin.com/in/sudip-bandopadhyay-18918a61'
)

# Summary
add_heading(doc, 'Professional Summary')
doc.add_paragraph(
    "Dynamic and accomplished Lead Data Engineer and Data Scientist with 14 years of experience delivering complex, large-scale data solutions across telecom and enterprise systems. "
    "Proficient in cloud-native data architectures (GCP & Azure), big data processing, and advanced analytics. Proven track record in optimizing data pipelines, building scalable platforms, "
    "and leading high-performing engineering teams. Recognized for reducing operational costs by up to 40% and enabling £7M+ business impact through strategic cloud migration, machine learning, and data warehousing solutions. "
    "Strong leadership in mentoring engineers and driving data strategy across cross-functional teams."
)

# Technical Skills
add_heading(doc, 'Technical Skills')
tech_skills = [
    "Cloud Platforms: Google Cloud Platform (GCP), Microsoft Azure, Oracle Cloud Infrastructure (OCI)",
    "Data Engineering & ETL Tools: Apache Airflow, Dataflow, Dataproc, Azure Data Factory, ODI, Informatica, DBT",
    "Programming: Python, SQL, PL/SQL, PySpark, UNIX Shell Scripting",
    "Data Warehousing: BigQuery, Snowflake, Oracle, SQL Server, PostgreSQL, Azure Synapse",
    "Streaming & Messaging: Apache Kafka, Google Pub/Sub",
    "Machine Learning & AI: TensorFlow, Keras, Scikit-learn, Vertex AI, Jupyter Notebooks, Pandas, NumPy",
    "Containerization & CICD: Docker, Kubernetes, GitHub, GitLab, Azure DevOps",
    "Reporting & Visualization: Looker, Matplotlib, Seaborn, Power BI",
    "Project & Code Management: Jira, Confluence, Git, SVN",
    "Architecture & Design: Star/Snowflake Schema, Data Lake, Data Mesh, Schema Design, Data Modeling",
]
add_bullets(doc, tech_skills)

# Certifications
add_heading(doc, 'Certifications')
certs = [
    "Microsoft Certified: Fabric Data Engineer Associate – 06/2025",
    "Google Cloud Certified: Professional Data Engineer – 04/2025",
    "Microsoft Certified: Azure Data Scientist Associate – 06/2024",
    "Oracle Cloud Infrastructure Foundations Associate – 07/2022",
    "Oracle Data Integrator 12c Implementation Specialist – 08/2019",
    "Oracle SOA Suite 11g Certified Implementation Specialist – 12/2016",
    "Oracle Database 11g: SQL Fundamentals – 12/2015",
]
add_bullets(doc, certs)

# Experience
add_heading(doc, 'Professional Experience')

doc.add_paragraph('Tata Consultancy Services — Lead Data Engineer | Data Scientist | Data Integration Specialist', style='Heading 2')
doc.add_paragraph('Kolkata, India / Whiteley, UK | 07/2014 – 05/2025')
tcs_bullets = [
    "Led enterprise-wide migration of on-premise Oracle data warehouse to GCP BigQuery, reducing operating costs by 40% and enabling £5M+ in cost savings.",
    "Built scalable ETL/ELT pipelines and automated orchestration using Airflow and Composer; optimized processing of petabyte-scale datasets.",
    "Developed ML models (customer propensity, fraud detection) using Vertex AI, Pandas, Scikit-learn; generated £2M+ in business impact.",
    "Created secure, production-ready data marts in BigQuery using DBT and optimized SQL logic for analytics use cases.",
    "Designed integration solutions using Oracle Fusion Middleware, ODI, SOA Suite for ERP system sync.",
    "Managed and mentored a 10-member engineering team with agile delivery practices.",
    "Collaborated with stakeholders for requirements gathering and roadmap planning.",
    "Received Best Performer of the Month (x2) and Best Team Award.",
]
add_bullets(doc, tcs_bullets)

doc.add_paragraph('Infosys Limited — ODI/ETL Developer', style='Heading 2')
doc.add_paragraph('Chennai, India | 08/2011 – 07/2014')
infosys_bullets = [
    "Developed robust ETL pipelines using Oracle Data Integrator (ODI) for telecom clients, improving transformation efficiency by 25%.",
    "Enhanced QA and release workflows, raising user satisfaction by 40%.",
    "Supported post-deployment integration and operations."
]
add_bullets(doc, infosys_bullets)

# Data Science Projects
add_heading(doc, 'Selected Data Science & AI Projects')
ml_projects = [
    "Customer Propensity Modeling: Supervised ML model boosted campaign targeting accuracy by 35% and delivered £2M+ in revenue.",
    "Fraud Detection Engine: Built classification model to identify anomalous transactions, reducing false positives by 45%.",
    "Chat Support Automation: NLP-based system cut ticket volume by 30% and improved customer support SLAs.",
    "Data Quality Dashboard: Real-time data health monitoring using Looker, GCP Functions, and Python.",
]
add_bullets(doc, ml_projects)

# Education
add_heading(doc, 'Education')
doc.add_paragraph('Bachelor of Technology – Electrical Engineering')
doc.add_paragraph('West Bengal University of Technology, India | 2011')

# Additional Info
add_heading(doc, 'Additional Information')
extra = [
    "Languages: English, Hindi, Bengali",
    "Awards: Best Performer of the Month (twice), Best Team Award",
    "Strengths: Strategic thinker, public speaker, fast learner, ready to travel globally",
]
add_bullets(doc, extra)

doc.save("Sudip_Bandopadhyay_Resume_3Page.docx")
print("Resume created successfully!")
